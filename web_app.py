# -*- coding: utf-8 -*-
"""
Flask Web Application for Compliance Checklist Automation
COMPLETE INTEGRATED VERSION with Smart Extraction Logic
"""

import os
import re
import json
import shutil
import time
import uuid
import traceback
import difflib
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
import datetime as _dt
import random
import threading
import logging

try:
    import requests
except Exception:
    requests = None

from flask import Flask, render_template, request, jsonify, send_file, session, abort
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.exceptions import HTTPException
import zipfile

# ===== Optional imports (graceful degradation) =====
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from PIL import Image, ImageFile, ImageOps, ImageEnhance, ImageFilter
    ImageFile.LOAD_TRUNCATED_IMAGES = True
    try:
        Image.MAX_IMAGE_PIXELS = 250_000_000
    except Exception:
        pass
except Exception:
    Image = None
    ImageOps = None
    ImageEnhance = None
    ImageFilter = None

try:
    import pytesseract
    if os.path.exists(r"C:\Program Files\Tesseract-OCR\tesseract.exe"):
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    elif os.path.exists("/usr/bin/tesseract"):
        pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"
except Exception:
    pytesseract = None

try:
    from tenacity import retry, stop_after_attempt, wait_exponential
except Exception:
    class MockRetry:
        def __init__(self, *args, **kwargs):
            pass
        def __call__(self, func):
            return func
    retry = MockRetry
    stop_after_attempt = lambda x: None
    wait_exponential = lambda **kwargs: None

try:
    from google import genai
except Exception:
    genai = None

try:
    import docx
    from docx.document import Document as _DocxDocument
except Exception as e:
    raise SystemExit("Missing dependency: python-docx. Install with: pip install python-docx") from e

try:
    from dateutil import parser as dateparser
except Exception:
    dateparser = None

# ===== Flask App Configuration =====
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'change-this-secret-key-in-production')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max upload

# ===== Always-JSON Error Handling (prevents HTML error pages) =====
DEBUG_ERRORS = os.getenv("DEBUG_ERRORS", "0").strip() == "1"

def _json_error(message: str, code: str = "ERROR", status: int = 400, details: Optional[dict] = None):
    payload = {
        'success': False,
        'error': {'code': code, 'message': message},
    }
    if DEBUG_ERRORS and details:
        payload['error']['details'] = details
    return jsonify(payload), status

@app.errorhandler(HTTPException)
def _handle_http_error(e: HTTPException):
    return _json_error(e.description, f"HTTP_{e.code}", e.code)

@app.errorhandler(Exception)
def _handle_unexpected_error(e: Exception):
    app.logger.exception("Unhandled error")
    details = {'type': type(e).__name__, 'trace': traceback.format_exc()}
    return _json_error("Request failed", "SERVER_ERROR", 500, details=details)

CORS(app)

@app.after_request
def _persist_ns_token(response):
    token = request.args.get('ns_token') or request.headers.get('X-NextStep-Token')
    if token:
        response.set_cookie('ns_token', token, httponly=True, samesite='Lax', secure=True, max_age=60 * 60 * 8, path='/')
    return response

# ── NextStep Auth ─────────────────────────────────────────────────────────────
BACKEND_BASE_URL = 'https://nextstep-backend-e75l.onrender.com'
BACKEND_VALIDATE_URL = f'{BACKEND_BASE_URL}/api/validate-session'
APP_DASHBOARD_URL = f'{BACKEND_BASE_URL}/dashboard'
APP_LOGIN_URL = f'{BACKEND_BASE_URL}/login'

def _get_ns_token(req):
    return (req.headers.get("X-NextStep-Token")
            or req.cookies.get("ns_token")
            or req.args.get("ns_token") or "")

def _validate_via_backend(token: str):
    if not token or requests is None:
        return None
    try:
        resp = requests.get(BACKEND_VALIDATE_URL, params={"token": token}, timeout=8)
        if resp.status_code != 200:
            return None
        data = resp.json()
        if not data.get('valid'):
            return None
        user = data.get('user') or {}
        tenant = data.get('tenant') or {}
        return {
            'user_id': user.get('id'),
            'tenant_id': tenant.get('id'),
            'role': user.get('role', 'admin'),
            'name': user.get('name'),
            'tenant_name': tenant.get('name'),
            'plan_name': tenant.get('plan_name') or '',
        }
    except Exception as e:
        app.logger.warning('[Auth backend] %s', e)
        return None

def _get_ctx():
    """Returns user context dict or None. Used in every route."""
    token = _get_ns_token(request)
    if not token:
        return None
    ctx = _validate_via_backend(token)
    if ctx:
        return ctx
    try:
        import db
        return db.validate_user_token(token)
    except Exception as e:
        app.logger.warning("[Auth db] %s", e)
        return None

def _require_auth():
    """Call at start of each route. Returns ctx dict or aborts 401."""
    from flask import abort
    ctx = _get_ctx()
    if not ctx:
        abort(401, f"Not authenticated. Please log in at {APP_LOGIN_URL}")
    return ctx

# ===== Folder Configuration =====
STORAGE_ROOT   = Path('/tmp/nextstep')
TEMPLATE_FOLDER = Path('templates_docx')
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'docx', 'txt'}
STORAGE_ROOT.mkdir(parents=True, exist_ok=True)
TEMPLATE_FOLDER.mkdir(exist_ok=True)

# Dynamic per-session paths — set in each request using tenant/user identity
def _upload_folder(tenant_id, user_id, session_id):
    p = STORAGE_ROOT / str(tenant_id) / str(user_id) / 'uploads' / session_id
    p.mkdir(parents=True, exist_ok=True)
    return p

def _output_folder(tenant_id, user_id, session_id):
    p = STORAGE_ROOT / str(tenant_id) / str(user_id) / 'outputs' / session_id
    p.mkdir(parents=True, exist_ok=True)
    return p


def _session_meta_dir(tenant_id, user_id, session_id):
    p = STORAGE_ROOT / str(tenant_id) / str(user_id) / 'sessions' / session_id
    p.mkdir(parents=True, exist_ok=True)
    return p


def _extracted_json_path(tenant_id, user_id, session_id):
    return _session_meta_dir(tenant_id, user_id, session_id) / 'extracted.json'


def _reviewed_json_path(tenant_id, user_id, session_id):
    return _session_meta_dir(tenant_id, user_id, session_id) / 'reviewed.json'

# Legacy fallback paths (used by cleanup only)
UPLOAD_FOLDER = STORAGE_ROOT / 'legacy_uploads'
OUTPUT_FOLDER = STORAGE_ROOT / 'legacy_outputs'

# ===== Gemini Configuration =====
GEMINI_MODEL_FAST = os.getenv("GEMINI_MODEL_FAST", "gemini-2.0-flash-001")
GEMINI_MODEL_STRONG = os.getenv("GEMINI_MODEL_STRONG", "gemini-2.5-pro")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
AI_CONF_THRESHOLD = float(os.getenv("AI_CONF_THRESHOLD", "0.65"))

# Initialize Gemini client
gemini_client = None
if genai and GEMINI_API_KEY:
    try:
        gemini_client = genai.Client(api_key=GEMINI_API_KEY)
    except Exception as e:
        print(f"Failed to initialize Gemini: {e}")

# ===== Constants =====
DOCUMENT_FOLDERS = [
    "APPLICATION FORM",
    "CV",
    "DBS",
    "NMC",
    "NI",
    "POA",
    "PASSPORT",
    "RTW",
    "TRAININGS",
    "IMAGE"
]

TEMPLATE_METADATA = [
    {'template_key': 'exemplar_profile', 'template_name': 'EXEMPLAR PROFILE', 'file_name': 'EXEMPLAR PROFILE.docx', 'active': True, 'order': 1},
    {'template_key': 'hc_one_profile', 'template_name': 'HC-One Profile', 'file_name': 'HC-One Profile.docx', 'active': True, 'order': 2},
    {'template_key': 'healthcare_homes_profile', 'template_name': 'HEALTHCARE HOMES PROFILE', 'file_name': 'HEALTHCARE HOMES PROFILE.docx', 'active': True, 'order': 3},
    {'template_key': 'horizon_care_profile', 'template_name': 'Horizon Care PROFILE', 'file_name': 'Horizon Care - PROFILE.docx', 'active': True, 'order': 4},
    {'template_key': 'iris_profile', 'template_name': 'IRIS PROFILE', 'file_name': 'IRIS PROFILE.docx', 'active': True, 'order': 5},
    {'template_key': 'lc_profile', 'template_name': 'LC Profile', 'file_name': 'LC Profile.docx', 'active': True, 'order': 6},
    {'template_key': 'mha_profile', 'template_name': 'MHA PROFILE', 'file_name': 'MHA PROFILE.docx', 'active': True, 'order': 7},
    {'template_key': 'neuven_new', 'template_name': 'neuven new', 'file_name': 'neuven new.docx', 'active': True, 'order': 8},
]

STANDARD_FIELDS = [
    "Candidate Name", "Title", "Forename(s)", "Surname", "Email",
    "Address", "Phone", "DOB", "Nationality", "NI Number",
    "Role", "NMC PIN", "DBS Number", "DBS Issue Date",
    "DBS Last Checked Date", "Training Date", "Training Expiry Date",
    "RTW Status", "Visa Expiry Date", "Visa Type", "Restriction", "Share Code",
    "Form Completed By", "Signature", "Position"
]

# Field extraction priority mapping
FIELD_PRIORITY = {
    "Candidate Name": ["DBS", "APPLICATION FORM", "CV", "PASSPORT"],
    "Title": ["APPLICATION FORM"],
    "Forename(s)": ["DBS", "APPLICATION FORM", "CV", "PASSPORT"],
    "Surname": ["DBS", "APPLICATION FORM", "CV", "PASSPORT"],
    "Email": ["APPLICATION FORM", "CV"],
    "Phone": ["APPLICATION FORM", "CV"],
    "DOB": ["DBS", "APPLICATION FORM", "PASSPORT"],
    "Address": ["POA", "APPLICATION FORM"],
    "Nationality": ["PASSPORT", "APPLICATION FORM"],
    "NI Number": ["NI", "APPLICATION FORM"],
    "Role": ["AUTO"],  # Special: NMC folder check
    "NMC PIN": ["NMC"],
    "DBS Number": ["DBS"],
    "DBS Issue Date": ["DBS"],
    "DBS Last Checked Date": ["DBS"],
    "Training Date": ["TRAININGS"],  # Special: earliest in 12 months
    "Training Expiry Date": ["CALCULATED"],  # Training Date + 12 months
    "RTW Status": ["RTW"],
    "Visa Expiry Date": ["RTW"],
    "Visa Type": ["RTW"],
    "Restriction": ["RTW"],
    "Share Code": ["RTW"]
}

SENSITIVE_FIELDS = {"NI Number"}
# --- Sensitive-field validation ---
NI_REGEX = re.compile(
    r"\b(?!BG)(?!GB)(?!NK)(?!KN)(?!TN)(?!NT)(?!ZZ)"
    r"[A-CEGHJ-PR-TW-Z]{2}\s?\d{2}\s?\d{2}\s?\d{2}\s?[A-D]\b",
    re.IGNORECASE
)

def validate_ni(value: str) -> str:
    v = (value or "").strip()
    if not v:
        return ""
    v_up = re.sub(r"\s+", " ", v).upper()
    v_nospace = v_up.replace(" ", "")
    if NI_REGEX.search(v_up) or NI_REGEX.search(v_nospace):
        return v_up
    return ""

def extract_print_date_from_text(text: str) -> str:
    """Extract Print Date from DBS check style text and treat as issue date."""
    if not text:
        return ""
    t = text.lower()

    # Common labels: print date / date printed / printed on
    patterns = [
        r"print\s*date\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"date\s*printed\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"printed\s*on\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
    ]
    for pat in patterns:
        m = re.search(pat, t, flags=re.I)
        if m:
            ds = m.group(1)
            try:
                if dateparser:
                    d = dateparser.parse(ds, dayfirst=True)
                    if d:
                        return d.date().strftime("%d/%m/%Y")
            except Exception:
                return ds

    # Fallback: find a date near the word 'print'
    idx = t.find('print')
    if idx != -1:
        window = t[max(0, idx-200): idx+200]
        m = re.search(r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", window)
        if m:
            ds = m.group(1)
            try:
                if dateparser:
                    d = dateparser.parse(ds, dayfirst=True)
                    if d:
                        return d.date().strftime("%d/%m/%Y")
            except Exception:
                return ds
    return ""

# --- Session cleanup (privacy + storage) ---
CLEANUP_AFTER_HOURS = int(os.getenv("CLEANUP_AFTER_HOURS", "24"))
DOWNLOAD_TTL_MINUTES = int(os.getenv('DOWNLOAD_TTL_MINUTES', '15'))
CHECKLIST_DISPATCH_POLL_SECONDS = float(os.getenv('CHECKLIST_DISPATCH_POLL_SECONDS', '2'))
CHECKLIST_STALE_MINUTES = int(os.getenv('CHECKLIST_STALE_MINUTES', '20'))

def cleanup_old_sessions() -> None:
    """Delete upload/output session folders older than CLEANUP_AFTER_HOURS."""
    cutoff = time.time() - (CLEANUP_AFTER_HOURS * 3600)

    def _cleanup_root(root: Path):
        if not root.exists():
            return
        for p in root.iterdir():
            try:
                if not p.is_dir():
                    continue
                mtime = p.stat().st_mtime
                if mtime < cutoff:
                    shutil.rmtree(p, ignore_errors=True)
            except Exception:
                continue

    _cleanup_root(UPLOAD_FOLDER)
    _cleanup_root(OUTPUT_FOLDER)

DEFAULT_PHONE = "0121 827 3666"

@dataclass
class FieldValue:
    value: str
    source: str
    confidence: float = 1.0
    ai_filled: bool = False

# ===== CORE EXTRACTION FUNCTIONS =====

def extract_text_pdf(pdf_path: Path) -> str:
    """Extract text from PDF using pdfplumber"""
    if not pdfplumber:
        return ""
    
    try:
        text_parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"PDF extraction error for {pdf_path}: {e}")
        return ""

def _ocr_score(text: str) -> int:
    """Heuristic score for OCR output quality."""
    if not text:
        return 0
    # Prefer longer, alphanumeric-heavy text; penalize too many symbols
    alnum = sum(ch.isalnum() for ch in text)
    bad = sum(ch in "|~`^_" for ch in text)
    return len(text) + 2 * alnum - 5 * bad

def _preprocess_for_ocr(img: 'Image.Image') -> 'Image.Image':
    """Lightweight preprocessing to improve OCR accuracy."""
    try:
        img = ImageOps.exif_transpose(img)  # fix orientation if present
    except Exception:
        pass

    # Convert to grayscale
    try:
        img = img.convert('L')
    except Exception:
        pass

    # Upscale small images (OCR works better)
    try:
        w, h = img.size
        if max(w, h) < 1400:
            scale = 1400 / max(w, h)
            img = img.resize((int(w * scale), int(h * scale)))
    except Exception:
        pass

    # Increase contrast and sharpen
    try:
        if ImageEnhance:
            img = ImageEnhance.Contrast(img).enhance(2.2)
            img = ImageEnhance.Sharpness(img).enhance(2.0)
    except Exception:
        pass

    # Mild denoise
    try:
        img = img.filter(ImageFilter.MedianFilter(size=3))
    except Exception:
        pass

    return img

def _ocr_image_best(img: 'Image.Image') -> str:
    """Run OCR with a few settings/orientations and keep best output."""
    if not pytesseract:
        return ""

    img = _preprocess_for_ocr(img)

    # Try common rotations (phone photos)
    rotations = [0, 90, 180, 270]
    psm_modes = [6, 11, 4]  # block text, sparse text, columns

    best_text = ""
    best_score = 0

    for rot in rotations:
        try:
            im2 = img.rotate(rot, expand=True) if rot else img
        except Exception:
            im2 = img

        for psm in psm_modes:
            cfg = f"--oem 3 --psm {psm}"
            try:
                t = pytesseract.image_to_string(im2, lang='eng', config=cfg) or ""
            except Exception:
                t = ""
            score = _ocr_score(t)
            if score > best_score:
                best_score = score
                best_text = t

    return (best_text or "").strip()

def extract_text_image(img_path: Path) -> str:
    """Extract text from image using OCR (enhanced)."""
    if not pytesseract or not Image:
        return ""

    try:
        ImageFile.LOAD_TRUNCATED_IMAGES = True
        img = Image.open(img_path)
        return _ocr_image_best(img)
    except Exception as e:
        print(f"Image OCR error for {img_path}: {e}")
        return ""


def extract_text_docx(docx_path: Path) -> str:
    """Extract text from DOCX"""
    try:
        doc = docx.Document(docx_path)
        text_parts = []
        
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        print(f"DOCX extraction error for {docx_path}: {e}")
        return ""

def gather_folder_text(folder: Path) -> Tuple[str, List[str], List[Path]]:
    """Gather all text from a specific folder.

    Returns:
      - combined text (with file markers),
      - processed file names,
      - processed file paths (for optional AI-vision fallback).
    """
    all_text: List[str] = []
    processed_files: List[str] = []
    processed_paths: List[Path] = []

    if not folder.exists():
        return "", [], []

    for file_path in folder.iterdir():
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        extracted_text = ""

        try:
            # Always track supported files for AI file-attachment fallback,
            # even if OCR/text-layer extraction returns empty text.
            if ext in {'.pdf', '.png', '.jpg', '.jpeg', '.docx', '.txt'}:
                processed_files.append(file_path.name)
                processed_paths.append(file_path)

            if ext == '.pdf':
                extracted_text = extract_text_pdf(file_path)
            elif ext in {'.png', '.jpg', '.jpeg'}:
                extracted_text = extract_text_image(file_path)
            elif ext == '.docx':
                extracted_text = extract_text_docx(file_path)
            elif ext == '.txt':
                # Be forgiving with encodings.
                extracted_text = file_path.read_text(encoding='utf-8', errors='ignore')

            if extracted_text:
                all_text.append(f"\n=== {file_path.name} ===\n{extracted_text}")

        except Exception as e:
            print(f"Error processing {file_path}: {e}")

    return "\n".join(all_text), processed_files, processed_paths


@retry(stop=stop_after_attempt(3), wait=wait_exponential(min=1, max=10))
def _gemini_generate(model: str, contents):
    """Small wrapper for Gemini calls (keeps retries consistent)."""
    return gemini_client.models.generate_content(model=model, contents=contents)

def _parse_json_response(response_text: str) -> Dict[str, Any]:
    """Parse JSON safely from Gemini text output."""
    if not response_text:
        return {}
    txt = response_text.strip()

    # Strip fenced blocks if present
    if txt.startswith('```'):
        txt = re.sub(r'^```(?:json)?\s*', '', txt, flags=re.I)
        txt = re.sub(r'\s*```\s*$', '', txt)

    # Quick repair: extract first {...} if model added extra text
    if not txt.startswith('{'):
        m = re.search(r'\{[\s\S]*\}', txt)
        if m:
            txt = m.group(0)

    try:
        return json.loads(txt)
    except Exception:
        return {}

def ai_extract_from_text(text: str, fields_to_extract: List[str]) -> Dict[str, FieldValue]:
    """Extract specific fields from text using Gemini AI.

    Hybrid model strategy (silent):
      - Primary: GEMINI_MODEL_FAST
      - Fallback: GEMINI_MODEL_STRONG
    """
    if not gemini_client or not text.strip():
        return {}

    prompt = f"""Extract the following information from the documents below. Return ONLY a JSON object with these exact fields:

{json.dumps(fields_to_extract, indent=2)}

Rules:
- Return ONLY valid JSON, nothing else
- Use exact field names as listed above
- If a field is not found, use empty string ""
- For dates, use DD/MM/YYYY format
- For NI Number, use format: XX 12 34 56 A
- Be accurate and extract exactly what you see

Documents:
{text[:15000]}

Return JSON only:"""

    models = [GEMINI_MODEL_FAST, GEMINI_MODEL_STRONG]

    last_err = None
    for model in models:
        try:
            resp = _gemini_generate(model=model, contents=prompt)
            data = _parse_json_response(getattr(resp, 'text', '') or '')
            if not isinstance(data, dict):
                data = {}

            result: Dict[str, FieldValue] = {}
            for field in fields_to_extract:
                value = data.get(field, "")
                if value and str(value).strip():
                    result[field] = FieldValue(
                        value=str(value).strip(),
                        source=f"AI (Gemini: {model})",
                        confidence=0.85,
                        ai_filled=True
                    )
            return result
        except Exception as e:
            last_err = e
            continue

    if last_err:
        print(f"AI extraction error (all models failed): {last_err}")
    return {}

def ai_extract_from_files(file_paths: List[Path], fields_to_extract: List[str]) -> Dict[str, FieldValue]:
    """AI vision-style fallback: attach files when OCR/text is weak.

    If the installed google-genai SDK doesn't support file parts in this environment,
    this function safely returns {}.
    """
    if not gemini_client or not file_paths:
        return {}

    try:
        from google.genai import types  # type: ignore
    except Exception:
        return {}

    # Attach a small number of files to control cost
    limited = file_paths[:4]
    parts = []
    for p in limited:
        try:
            b = p.read_bytes()
            mime = "application/octet-stream"
            ext = p.suffix.lower()
            if ext in ['.jpg', '.jpeg']:
                mime = 'image/jpeg'
            elif ext == '.png':
                mime = 'image/png'
            elif ext == '.pdf':
                mime = 'application/pdf'
            elif ext == '.docx':
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            parts.append(types.Part.from_bytes(data=b, mime_type=mime))
        except Exception:
            continue

    if not parts:
        return {}

    prompt = f"""Extract the following information from the attached documents. Return ONLY a JSON object with these exact fields:

{json.dumps(fields_to_extract, indent=2)}

Rules:
- Return ONLY valid JSON, nothing else
- Use exact field names as listed above
- If a field is not found, use empty string ""
- For dates, use DD/MM/YYYY format
- Be accurate and extract exactly what you see

Return JSON only:"""

    models = [GEMINI_MODEL_FAST, GEMINI_MODEL_STRONG]
    last_err = None
    for model in models:
        try:
            resp = _gemini_generate(model=model, contents=[prompt, *parts])
            data = _parse_json_response(getattr(resp, 'text', '') or '')
            if not isinstance(data, dict):
                data = {}

            result: Dict[str, FieldValue] = {}
            for field in fields_to_extract:
                value = data.get(field, "")
                if value and str(value).strip():
                    result[field] = FieldValue(
                        value=str(value).strip(),
                        source=f"AI (Gemini attach: {model})",
                        confidence=0.85,
                        ai_filled=True
                    )
            return result
        except Exception as e:
            last_err = e
            continue

    if last_err:
        print(f"AI file-attach extraction error: {last_err}")
    return {}


# ===== Gemini Vision Fallback (PDF->Images via PyMuPDF, no OCR) =====

def pdf_to_images_bytes(pdf_path: Path, max_pages: int = 5) -> List[bytes]:
    """Convert PDF pages to PNG bytes using PyMuPDF. Works for scanned PDFs."""
    if not fitz:
        return []
    images: List[bytes] = []
    try:
        doc = fitz.open(str(pdf_path))
        pages_to_render = min(doc.page_count, max_pages)
        mat = fitz.Matrix(2, 2)
        for i in range(pages_to_render):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("png"))
        doc.close()
    except Exception as e:
        print(f"PyMuPDF render error for {pdf_path}: {e}")
    return images

def ai_extract_from_vision_files(file_paths: List[Path], fields_to_extract: List[str]) -> Dict[str, FieldValue]:
    """Vision fallback: attach images (and rendered PDF pages) to Gemini.

    - Keeps your label-based rule (no guessing; if unsure -> blank).
    - Uses same hybrid model strategy (FAST then STRONG).
    - Fills ONLY provided fields.
    """
    if not gemini_client or not file_paths:
        return {}

    try:
        from google.genai import types  # type: ignore
    except Exception:
        return {}

    # Build image parts (PDF->images; images direct). Cap to control cost.
    parts: List[Any] = []
    for p in file_paths:
        ext = p.suffix.lower()
        try:
            if ext in ['.jpg', '.jpeg', '.png']:
                b = p.read_bytes()
                if not b:
                    continue
                mime = 'image/jpeg' if ext in ['.jpg', '.jpeg'] else 'image/png'
                parts.append(types.Part.from_bytes(data=b, mime_type=mime))
            elif ext == '.pdf':
                for img in pdf_to_images_bytes(p, max_pages=5):
                    parts.append(types.Part.from_bytes(data=img, mime_type='image/png'))
        except Exception:
            continue
        if len(parts) >= 12:
            break

    if not parts:
        return {}

    # Label synonyms (still label-anchored, not regex guessing)
    label_aliases = {
        "Candidate Name": ["Name", "Full Name", "Applicant Name", "Candidate"],
        "DOB": ["DOB", "Date of Birth", "Birth Date"],
        "DBS Number": ["DBS Number", "DBS No", "Certificate Number", "Certificate No"],
        "DBS Issue Date": ["DBS Issue Date", "Issue Date", "Date of Issue", "Print Date", "Printed on", "Date Printed"],
        "NI Number": ["NI Number", "National Insurance Number", "NINO"],
        "NMC PIN": ["NMC PIN", "PIN", "NMC Pin Number"],
        "Share Code": ["Share Code"],
        "Visa Expiry Date": ["Visa Expiry Date", "Expiry Date"],
        "Visa Type": ["Visa Type", "Type of visa"],
        "Restriction": ["Restriction"],
        "RTW Status": ["Right to Work", "RTW Status"],
        "Address": ["Address", "Home Address"],
        "Phone": ["Phone", "Mobile", "Telephone"],
        "Email": ["Email", "E-mail"],
        "Nationality": ["Nationality"],
        "Surname": ["Surname", "Last Name", "Family Name"],
        "Forename(s)": ["Forename", "First Name", "Given Name"],
    }
    alias_lines = []
    for f in fields_to_extract:
        alias_lines.append(f"- {f}: labels may appear as {', '.join([repr(a) for a in label_aliases.get(f, [])])}")

    prompt = f"""Extract the following information from the attached document images. Return ONLY a JSON object with these exact fields:

{json.dumps(fields_to_extract, indent=2)}

RULES:
- Label-based extraction only: extract a value ONLY if it is clearly present near an explicit label (or a label synonym listed below).
- Do NOT guess or infer. If unsure or not visible, return empty string "".
- Do NOT invent sensitive identifiers (NI/DBS/NMC/Passport/Share Code). If not clearly visible, return "".
- Candidate Name must be the PERSON/APPLICANT name, not employer/company.

Label synonyms you may treat as equivalent:
{chr(10).join(alias_lines)}

DBS rules:
- "Printed on"/"Print Date" is the DBS Issue Date.
- "Certificate Number" is the DBS Number.

Return JSON only."""

    models = [GEMINI_MODEL_FAST, GEMINI_MODEL_STRONG]
    last_err = None
    for model in models:
        try:
            # Prefer Content(role='user', parts=[text, images...]) when available
            try:
                contents = [types.Content(role="user", parts=[types.Part.from_text(prompt)] + parts)]
            except Exception:
                contents = [prompt, *parts]
            resp = _gemini_generate(model=model, contents=contents)
            data = _parse_json_response(getattr(resp, 'text', '') or '')
            if not isinstance(data, dict):
                data = {}

            result: Dict[str, FieldValue] = {}
            for field in fields_to_extract:
                value = data.get(field, "")
                if value and str(value).strip():
                    result[field] = FieldValue(
                        value=str(value).strip(),
                        source=f"AI (Gemini Vision: {model})",
                        confidence=0.85,
                        ai_filled=True
                    )
            return result
        except Exception as e:
            last_err = e
            continue

    if last_err:
        print(f"AI vision extraction error: {last_err}")
    return {}

def detect_role_from_nmc(nmc_folder: Path) -> str:
    """Detect role based on NMC folder contents"""
    if nmc_folder.exists() and any(nmc_folder.iterdir()):
        return "RGN"
    return "HCA"

def extract_training_date(trainings_folder: Path) -> Tuple[str, str, float]:
    """Extract training date.

    Rule:
    - Find the earliest date within last 12 months across training documents.
    - If none found, generate a random training date (last 3 months).
    Also returns the exact source file that contained the chosen date.
    """
    if not trainings_folder.exists():
        return generate_random_training_date()

    # Parse per-file so we can report correct source
    date_pattern = re.compile(r'\b(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b')
    today = _dt.date.today()
    twelve_months_ago = today - _dt.timedelta(days=365)

    best_date: Optional[_dt.date] = None
    best_file: Optional[str] = None

    _, _, paths = gather_folder_text(trainings_folder)

    for p in paths:
        ext = p.suffix.lower()
        if ext == '.pdf':
            t = extract_text_pdf(p)
        elif ext in ['.png', '.jpg', '.jpeg']:
            t = extract_text_image(p)
        elif ext == '.docx':
            t = extract_text_docx(p)
        else:
            t = ''

        if not t:
            continue

        for date_str in date_pattern.findall(t):
            try:
                if not dateparser:
                    continue
                parsed = dateparser.parse(date_str, dayfirst=True)
                if not parsed:
                    continue
                d = parsed.date()
                if twelve_months_ago <= d <= today:
                    if best_date is None or d < best_date:
                        best_date = d
                        best_file = p.name
            except Exception:
                continue

    if best_date:
        src = f"TRAININGS/{best_file}" if best_file else "TRAININGS/document"
        return best_date.strftime("%d/%m/%Y"), src, 0.90

    return generate_random_training_date()


def generate_random_training_date() -> Tuple[str, str, float]:
    """Generate random training date from last 3 months"""
    today = _dt.date.today()
    days_ago = random.randint(1, 90)
    random_date = today - _dt.timedelta(days=days_ago)
    return random_date.strftime("%d/%m/%Y"), "Random (no valid training in last 12 months)", 0.50

def calculate_training_expiry(training_date: str) -> str:
    """Calculate training expiry as training date + 12 months"""
    try:
        if dateparser:
            parsed = dateparser.parse(training_date, dayfirst=True)
            if parsed:
                expiry = parsed.date() + _dt.timedelta(days=365)
                return expiry.strftime("%d/%m/%Y")
    except:
        pass
    return ""

def check_uk_passport(rtw_folder: Path, nationality: str) -> bool:
    """Check if UK Passport detected in RTW folder or nationality"""
    if "british" in nationality.lower():
        return True
    
    if rtw_folder.exists():
        text, _, _ = gather_folder_text(rtw_folder)
        if "uk passport" in text.lower() or "british passport" in text.lower():
            return True
    
    return False

def auto_derive_names(extracted: Dict[str, FieldValue]) -> Dict[str, FieldValue]:
    """Auto-derive Candidate Name from Forename + Surname and vice versa"""
    candidate_name = extracted.get("Candidate Name", FieldValue("", ""))
    forename = extracted.get("Forename(s)", FieldValue("", ""))
    surname = extracted.get("Surname", FieldValue("", ""))
    
    # If we have Candidate Name but not Forename/Surname
    if candidate_name.value and (not forename.value or not surname.value):
        parts = candidate_name.value.split()
        if len(parts) >= 2:
            if not surname.value:
                extracted["Surname"] = FieldValue(
                    parts[-1],
                    f"Derived from Candidate Name ({candidate_name.source})",
                    candidate_name.confidence * 0.9,
                    True
                )
            if not forename.value:
                extracted["Forename(s)"] = FieldValue(
                    " ".join(parts[:-1]),
                    f"Derived from Candidate Name ({candidate_name.source})",
                    candidate_name.confidence * 0.9,
                    True
                )
    
    # If we have Forename + Surname but not Candidate Name
    elif forename.value and surname.value and not candidate_name.value:
        extracted["Candidate Name"] = FieldValue(
            f"{forename.value} {surname.value}",
            f"Derived from Forename + Surname",
            min(forename.confidence, surname.confidence) * 0.95,
            True
        )
    
    return extracted

def smart_extract_with_priority(session_folder: Path) -> Dict[str, FieldValue]:
    """Extract fields with folder priority logic (with safe fallbacks)."""
    extracted: Dict[str, FieldValue] = {}
    folder_texts: Dict[str, str] = {}
    folder_files: Dict[str, List[str]] = {}
    folder_paths: Dict[str, List[Path]] = {}

    # Gather text + assets from each folder
    for folder_name in DOCUMENT_FOLDERS:
        folder_path = session_folder / folder_name
        text_blob, files, paths = gather_folder_text(folder_path)
        folder_texts[folder_name] = text_blob
        folder_files[folder_name] = files
        folder_paths[folder_name] = paths

    # Special: Role detection from NMC folder
    role = detect_role_from_nmc(session_folder / "NMC")
    extracted["Role"] = FieldValue(role, "AUTO (NMC folder detection)", 1.0, False)

    # Special: Training date logic
    training_date, training_source, training_conf = extract_training_date(session_folder / "TRAININGS")
    extracted["Training Date"] = FieldValue(training_date, training_source, training_conf, True)

    # Special: Training expiry = Training date + 12 months
    training_expiry = calculate_training_expiry(training_date)
    if training_expiry:
        extracted["Training Expiry Date"] = FieldValue(
            training_expiry,
            "Calculated (Training Date + 12 months)",
            training_conf,
            True
        )

    # Extract fields based on priority
    for field, priority_folders in FIELD_PRIORITY.items():
        if field in ["Role", "Training Date", "Training Expiry Date"]:
            continue  # Already handled

        # NMC PIN only relevant if role is RGN
        if field == "NMC PIN":
            if role == "RGN" and folder_texts.get("NMC"):
                ai_result = ai_extract_from_text(folder_texts["NMC"], ["NMC PIN"])
                if "NMC PIN" in ai_result:
                    extracted["NMC PIN"] = ai_result["NMC PIN"]
                    if folder_files.get('NMC'):
                        extracted["NMC PIN"].source = f"NMC/{folder_files.get('NMC')[0]}"
            else:
                extracted["NMC PIN"] = FieldValue("NA", "Role is HCA", 1.0, False)
            continue

        # Try folders in priority order
        for folder in priority_folders:
            if folder in ("AUTO", "CALCULATED"):
                continue

            # 1) Text-based AI extraction when we have text
            if folder_texts.get(folder):
                ai_result = ai_extract_from_text(folder_texts[folder], [field])
                if field in ai_result:
                    extracted[field] = ai_result[field]
                    if folder_files.get(folder):
                        extracted[field].source = f"{folder}/{folder_files.get(folder)[0]}"
                    else:
                        extracted[field].source = f"{folder}"
                    break

            # 2) Vision-style fallback if folder has files but text is weak/empty
            if folder_paths.get(folder):
                # only attempt if not already extracted and we have almost no text
                if (not folder_texts.get(folder)) or (len(folder_texts.get(folder) or "") < 80):
                    ai_result = ai_extract_from_files(folder_paths[folder], [field])
                    if field in ai_result:
                        extracted[field] = ai_result[field]
                        extracted[field].source = f"{folder} (attachment fallback)"
                        break

                    # 3) NEW: Gemini Vision fallback (PDF->images / image attachments) if still blank
                    vision_result = ai_extract_from_vision_files(folder_paths[folder], [field])
                    if field in vision_result:
                        extracted[field] = vision_result[field]
                        extracted[field].source = f"{folder} (vision fallback)"
                        break

    # DBS Issue Date: if not extracted from certificate image, use Print Date from DBS check PDF
    if not (extracted.get("DBS Issue Date") and extracted["DBS Issue Date"].value.strip()):
        dbs_text = folder_texts.get("DBS", "")
        issue = extract_print_date_from_text(dbs_text)
        if issue:
            extracted["DBS Issue Date"] = FieldValue(
                issue,
                "DBS (Print Date used as Issue Date)",
                0.80,
                True
            )

    # Auto-derive names
    extracted = auto_derive_names(extracted)

    # Check UK Passport for RTW fields
    nationality = extracted.get("Nationality", FieldValue("", "")).value
    is_uk_passport = check_uk_passport(session_folder / "RTW", nationality)

    if is_uk_passport:
        for f in ["RTW Status", "Visa Expiry Date", "Visa Type", "Restriction", "Share Code"]:
            extracted[f] = FieldValue("UK Passport", "UK Passport detected", 1.0, False)

    # Validate sensitive fields (never keep invented identifiers)
    for f in SENSITIVE_FIELDS:
        if f in extracted:
            if f == "NI Number":
                cleaned = validate_ni(extracted[f].value)
                if not cleaned:
                    extracted[f] = FieldValue("", "Blanked (invalid NI format)", 0.0, False)
                else:
                    extracted[f].value = cleaned

    return extracted


def replace_placeholders_in_paragraph(paragraph, replacements: Dict[str, str], fuzzy_cutoff: float = 0.75) -> bool:
    """Replace placeholders in a paragraph deterministically (run-safe), with fuzzy fallback."""
    original = paragraph.text or ""
    if not original:
        return False

    def _norm_key(k: str) -> str:
        k = (k or "").replace("’", "'")
        k = k.strip().lower()
        return re.sub(r"[^a-z0-9]+", "", k)

    norm_map = {_norm_key(k): (v or "") for k, v in replacements.items()}
    norm_keys = list(norm_map.keys())

    def _lookup(inner: str) -> Optional[str]:
        nk = _norm_key(inner)
        if nk in norm_map:
            return norm_map[nk]
        # aggressive fuzzy match (approved)
        if nk and norm_keys:
            best = difflib.get_close_matches(nk, norm_keys, n=1, cutoff=fuzzy_cutoff)
            if best:
                return norm_map.get(best[0])
        return None

    text = original.replace("{[", "{{")
    changed = False

    # handle nested placeholders like {{{{A}}/B}} style
    nested_re = re.compile(r"\{\{\{\{\s*([^{}]+?)\s*\}\}\s*/\s*[^{}]+?\s*\}\}\s*", re.I)

    def _nested_sub(m):
        inner = m.group(1)
        val = _lookup(inner)
        nonlocal changed
        changed = True
        return (val or "") if val is not None else ""

    if nested_re.search(text):
        text = nested_re.sub(_nested_sub, text)

    token_re = re.compile(r"\{\{([^{}]+)\}\}")
    for _ in range(5):
        made = False

        def _sub(m):
            inner = m.group(1)
            val = _lookup(inner)
            if val is None:
                return m.group(0)
            nonlocal made
            made = True
            return val

        new_text = token_re.sub(_sub, text)
        if new_text != text:
            text = new_text
        if not made:
            break
        changed = True

    if "{{" in text and "}}" in text:
        text2 = text.replace("{{", "").replace("}}", "")
        if text2 != text:
            text = text2
            changed = True

    if not changed:
        return False

    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
    return True

def replace_placeholders_in_table(table, replacements: Dict[str, str]) -> int:
    changed = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if replace_placeholders_in_paragraph(p, replacements):
                    changed += 1
            for t in cell.tables:
                changed += replace_placeholders_in_table(t, replacements)
    return changed

def _find_unfilled_placeholders_docx(d: '_DocxDocument') -> List[str]:
    token_re = re.compile(r"\{\{[^{}]+\}\}")
    leftovers = set()

    def scan_paras(paras):
        for p in paras:
            for tok in token_re.findall(p.text or ""):
                leftovers.add(tok)

    scan_paras(d.paragraphs)

    def scan_table(table):
        for row in table.rows:
            for cell in row.cells:
                scan_paras(cell.paragraphs)
                for t in cell.tables:
                    scan_table(t)

    for t in d.tables:
        scan_table(t)

    try:
        for section in d.sections:
            scan_paras(section.header.paragraphs)
            for t in section.header.tables:
                scan_table(t)
            scan_paras(section.footer.paragraphs)
            for t in section.footer.tables:
                scan_table(t)
    except Exception:
        pass

    return sorted(leftovers)

def fill_docx_template(template_path: Path, output_path: Path, replacements: Dict[str, str]) -> Tuple[int, List[str]]:
    """Fill DOCX template by replacing placeholders robustly (run-safe + tables + header/footer)."""
    warnings: List[str] = []
    d = docx.Document(str(template_path))
    changed = 0

    for p in d.paragraphs:
        if replace_placeholders_in_paragraph(p, replacements):
            changed += 1
    for t in d.tables:
        changed += replace_placeholders_in_table(t, replacements)

    try:
        for section in d.sections:
            header = section.header
            for p in header.paragraphs:
                if replace_placeholders_in_paragraph(p, replacements):
                    changed += 1
            for t in header.tables:
                changed += replace_placeholders_in_table(t, replacements)

            footer = section.footer
            for p in footer.paragraphs:
                if replace_placeholders_in_paragraph(p, replacements):
                    changed += 1
            for t in footer.tables:
                changed += replace_placeholders_in_table(t, replacements)
    except Exception as e:
        warnings.append(f"Header/footer replace skipped: {e}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(output_path))

    # Unfilled placeholder report
    try:
        d2 = docx.Document(str(output_path))
        leftovers = _find_unfilled_placeholders_docx(d2)
        if leftovers:
            report = output_path.parent / "unfilled_placeholders.txt"
            with report.open('a', encoding='utf-8') as f:
                f.write(f"\n=== {output_path.name} ===\n")
                for tok in leftovers:
                    f.write(tok + "\n")
    except Exception as e:
        warnings.append(f"Unfilled placeholder scan skipped: {e}")

    return changed, warnings

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def list_template_metadata() -> List[dict]:
    metas = []
    for item in sorted(TEMPLATE_METADATA, key=lambda x: x.get('order', 999)):
        p = TEMPLATE_FOLDER / item['file_name']
        if p.exists() and item.get('active', True):
            metas.append(dict(item))
    return metas


def get_template_by_key(template_key: str) -> Optional[dict]:
    for item in list_template_metadata():
        if item['template_key'] == template_key:
            return item
    return None


def list_templates() -> List[str]:
    return [t['template_name'] for t in list_template_metadata()]

# ===== Checklist worker / queue =====
CHECKLIST_GLOBAL_CONCURRENCY = max(1, int(os.getenv('CHECKLIST_MAX_CONCURRENT_PARENTS', '3')))
_dispatcher_started = False
_dispatcher_lock = threading.Lock()


def _safe_filename_part(value: str, fallback: str = 'Candidate') -> str:
    value = (value or '').strip() or fallback
    value = re.sub(r'[\/:*?"<>|]+', ' ', value)
    value = re.sub(r'\s+', '_', value).strip('._')
    return value[:80] or fallback


def _today_str() -> str:
    return _dt.date.today().strftime('%Y%m%d')


def _build_output_filename(candidate_name: str, template_name: str, output_folder: Path) -> str:
    base = f"{_safe_filename_part(candidate_name)}_{_safe_filename_part(template_name, 'Checklist')}_{_today_str()}"
    candidate = base + '.docx'
    n = 2
    while (output_folder / candidate).exists():
        candidate = f"{base}_v{n}.docx"
        n += 1
    return candidate


def _build_zip_filename(candidate_name: str) -> str:
    return f"{_safe_filename_part(candidate_name)}_Checklist_{_today_str()}.zip"


def _meaningful_extraction_success(extracted: Dict[str, FieldValue]) -> bool:
    meaningful = 0
    for _, fv in (extracted or {}).items():
        val = (fv.value or '').strip() if hasattr(fv, 'value') else str(fv or '').strip()
        if val and val.upper() not in {'NA', 'UK PASSPORT'}:
            meaningful += 1
    return meaningful >= 5


def _write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2), encoding='utf-8')


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding='utf-8'))



def _cleanup_downloaded_outputs():
    import db
    ttl_seconds = max(60, DOWNLOAD_TTL_MINUTES * 60)
    now = _dt.datetime.utcnow()
    outputs = []
    try:
        session = db.get_session()
        if not session:
            return
        try:
            rows = session.execute(__import__('sqlalchemy').text("SELECT output_id, output_path, session_id, tenant_id, user_id, first_downloaded_at, last_downloaded_at FROM checklist_outputs WHERE first_downloaded_at IS NOT NULL")).mappings().all()
            outputs = [dict(r) for r in rows]
        finally:
            session.close()
    except Exception:
        return
    for out in outputs:
        ts = out.get('last_downloaded_at') or out.get('first_downloaded_at')
        if not ts:
            continue
        try:
            if isinstance(ts, str):
                ts_dt = _dt.datetime.fromisoformat(ts.replace('Z', '+00:00').replace(' ', 'T'))
                if ts_dt.tzinfo is not None:
                    ts_dt = ts_dt.astimezone(_dt.timezone.utc).replace(tzinfo=None)
            else:
                ts_dt = ts
        except Exception:
            continue
        if (now - ts_dt).total_seconds() < ttl_seconds:
            continue
        try:
            fp = Path(out['output_path'])
            if fp.exists():
                fp.unlink(missing_ok=True)
            # remove zip too if exists in same session folder
            zf = fp.parent / _build_zip_filename((_read_json(_reviewed_json_path(out['tenant_id'], out['user_id'], out['session_id'])) or {}).get('Candidate Name') or 'Candidate')
            if zf.exists():
                zf.unlink(missing_ok=True)
        except Exception:
            pass

def _run_checklist_extraction(job: dict):
    import db
    tenant_id = int(job['tenant_id']); user_id = int(job['user_id'])
    session_id = job['session_id']; job_id = job['job_id']
    try:
        upload_folder = _upload_folder(tenant_id, user_id, session_id)
        db.update_checklist_job(job_id, tenant_id, user_id, current_step='Checking priority folders and extracting fields')
        extracted = smart_extract_with_priority(upload_folder)
        if not _meaningful_extraction_success(extracted):
            raise ValueError('No usable structured data could be extracted')
        extracted_data = {field: {'value': fv.value, 'source': fv.source, 'confidence': fv.confidence} for field, fv in extracted.items()}
        extracted_path = _extracted_json_path(tenant_id, user_id, session_id)
        _write_json(extracted_path, extracted_data)
        usage_id = None
        if not job.get('extraction_token_charged'):
            legacy_job_id = db.create_legacy_job_record(tenant_id=tenant_id, user_id=user_id, total_items=1, status='running')
            usage_id = db.record_usage(tenant_id=tenant_id, user_id=user_id, db_job_id=legacy_job_id, successful_outputs=1)
            if legacy_job_id:
                db.update_legacy_job_status(db_job_id=legacy_job_id, status='completed', successful_items=1, failed_items=0)
        db.update_checklist_job(job_id, tenant_id, user_id, status='extract_completed', current_step='Extraction completed', error_message=None, extracted_data_path=str(extracted_path), extraction_token_charged=True, extraction_token_charge_id=usage_id, extract_completed_at=_dt.datetime.utcnow())
    except Exception as e:
        app.logger.exception('[Checklist] extraction failed for %s', job_id)
        db.update_checklist_job(job_id, tenant_id, user_id, status='extract_failed', current_step='Extraction failed', error_message=str(e))


def _dispatcher_cycle():
    import db
    db.requeue_stale_running_jobs(minutes=30)
    total_running = db.count_total_running_jobs()
    if total_running >= CHECKLIST_GLOBAL_CONCURRENCY:
        return
    for job in db.list_queued_jobs(limit=50):
        if total_running >= CHECKLIST_GLOBAL_CONCURRENCY:
            break
        plan_limit = 3 if 'standard' in str(job.get('plan_type') or '').lower() or 'premium' in str(job.get('plan_type') or '').lower() or 'pro' in str(job.get('plan_type') or '').lower() else 1
        if db.count_running_jobs_for_user(int(job['tenant_id']), int(job['user_id'])) >= plan_limit:
            continue
        if db.claim_queued_job(job['job_id']):
            total_running += 1
            threading.Thread(target=_run_checklist_extraction, args=(job,), daemon=True, name=f"checklist-{job['job_id'][:8]}").start()


def _dispatcher_loop():
    while True:
        try:
            _dispatcher_cycle()
        except Exception:
            app.logger.exception('[Checklist] dispatcher loop error')
        time.sleep(CHECKLIST_DISPATCH_POLL_SECONDS if CHECKLIST_DISPATCH_POLL_SECONDS > 1 else 2)


def _ensure_dispatcher_started():
    global _dispatcher_started
    with _dispatcher_lock:
        if _dispatcher_started:
            return
        try:
            import db
            db.ensure_schema()
        except Exception:
            app.logger.exception('[Checklist] ensure_schema failed')
        threading.Thread(target=_dispatcher_loop, daemon=True, name='checklist-dispatcher').start()
        _dispatcher_started = True


# ===== FLASK ROUTES =====

@app.before_request
def _start_bg_worker():
    _ensure_dispatcher_started()


@app.route('/')
def index():
    token = _get_ns_token(request)
    ctx = _get_ctx()
    response = render_template('index.html', folders=DOCUMENT_FOLDERS, templates=list_templates(), fields=STANDARD_FIELDS, auth_required=(not bool(ctx) and not bool(token)), dashboard_url=APP_DASHBOARD_URL, login_url=APP_LOGIN_URL, download_ttl_minutes=DOWNLOAD_TTL_MINUTES, has_token=bool(token))
    from flask import make_response
    resp = make_response(response)
    if token and ctx:
        resp.set_cookie('ns_token', token, httponly=True, samesite='Lax', secure=True, max_age=60 * 60 * 8, path='/')
    return resp


@app.route('/templates', methods=['GET'])
def templates_api():
    _require_auth()
    return jsonify({'success': True, 'templates': list_template_metadata()})


@app.route('/upload', methods=['POST'])
def upload_files():
    ctx = _require_auth()
    tenant_id = ctx['tenant_id']; user_id = ctx['user_id']
    try:
        cleanup_old_sessions()
        session_id = str(uuid.uuid4())
        job_id = f"chk_{uuid.uuid4().hex}"
        session_folder = _upload_folder(tenant_id, user_id, session_id)
        upload_stats = {}
        total_files = 0
        for folder_name in DOCUMENT_FOLDERS:
            folder_path = session_folder / folder_name
            folder_path.mkdir(parents=True, exist_ok=True)
            field_name = f'files_{folder_name.replace(" ", "_")}'
            files = [f for f in request.files.getlist(field_name) if f and (f.filename or '').strip()]
            allowed_limit = None if folder_name == 'TRAININGS' else 5
            if allowed_limit is not None and len(files) > allowed_limit:
                return jsonify({'success': False, 'error_message': f'{folder_name} allows maximum {allowed_limit} files. You selected {len(files)}.'}), 400
            count = 0
            used_names = set()
            for file in files:
                if not allowed_file(file.filename):
                    continue
                filename = secure_filename(file.filename)
                if not filename:
                    continue
                stem = Path(filename).stem
                suffix = Path(filename).suffix
                candidate = filename
                n = 2
                while candidate.lower() in used_names or (folder_path / candidate).exists():
                    candidate = f"{stem}_{n}{suffix}"
                    n += 1
                used_names.add(candidate.lower())
                file.save(folder_path / candidate)
                count += 1
                total_files += 1
            upload_stats[folder_name] = count
        if total_files == 0:
            return jsonify({'success': False, 'error_message': 'Please upload at least one supported file.'}), 400
        import db
        db.create_checklist_job(job_id=job_id, session_id=session_id, tenant_id=tenant_id, user_id=user_id,
                                plan_type=ctx.get('plan_name') or '', upload_path=str(session_folder), uploaded_file_count=total_files)
        session['session_id'] = session_id
        session['tenant_id'] = tenant_id
        session['user_id'] = user_id
        return jsonify({'success': True, 'job_id': job_id, 'session_id': session_id, 'status': 'uploaded', 'uploaded_file_count': total_files, 'upload_stats': upload_stats, 'message': 'Files uploaded successfully.'})
    except Exception as e:
        app.logger.exception('upload failed')
        return jsonify({'success': False, 'error_message': str(e)}), 500


@app.route('/extract', methods=['POST'])
def extract_data():
    ctx = _require_auth()
    tenant_id = ctx['tenant_id']; user_id = ctx['user_id']
    data = request.get_json(silent=True) or {}
    job_id = data.get('job_id')
    session_id = data.get('session_id')
    if not job_id or not session_id:
        return jsonify({'success': False, 'error_message': 'Missing job_id or session_id'}), 400
    try:
        import db
        tokens = db.get_tenant_tokens_remaining(tenant_id)
        if tokens == 0:
            return jsonify({'success': False, 'error_message': 'No tokens remaining. Please contact NextStep to top up.'}), 402
        job = db.get_checklist_job(job_id, tenant_id, user_id)
        if not job or job.get('session_id') != session_id:
            return jsonify({'success': False, 'error_message': 'Job not found'}), 404
        if not db.queue_checklist_job(job_id, session_id, tenant_id, user_id):
            return jsonify({'success': False, 'error_message': 'Job cannot be queued from its current state.'}), 409
        return jsonify({'success': True, 'job_id': job_id, 'session_id': session_id, 'status': 'extract_queued', 'message': 'Extraction queued successfully.'})
    except Exception as e:
        app.logger.exception('extract queue failed')
        return jsonify({'success': False, 'error_message': str(e)}), 500


@app.route('/extract/status/<job_id>', methods=['GET'])
def extract_status(job_id):
    ctx = _require_auth()
    import db
    job = db.get_checklist_job(job_id, ctx['tenant_id'], ctx['user_id'])
    if not job:
        return jsonify({'success': False, 'error_message': 'Job not found'}), 404
    success = job['status'] != 'extract_failed'
    return jsonify({
        'success': success,
        'job_id': job['job_id'],
        'session_id': job['session_id'],
        'status': job['status'],
        'current_step': job.get('current_step'),
        'error_message': job.get('error_message'),
        'review_ready': job['status'] == 'extract_completed',
    }), (200 if success else 500)


@app.route('/extracted/<session_id>', methods=['GET'])
def get_extracted(session_id):
    ctx = _require_auth()
    import db
    job = db.get_job_by_session(session_id, ctx['tenant_id'], ctx['user_id'])
    if not job:
        return jsonify({'success': False, 'error_message': 'Session not found'}), 404
    if job['status'] not in ('extract_completed', 'review_ready', 'generating', 'generated'):
        return jsonify({'success': False, 'error_message': 'Extraction not completed yet'}), 409
    path = Path(job.get('reviewed_data_path') or job.get('extracted_data_path') or _extracted_json_path(ctx['tenant_id'], ctx['user_id'], session_id))
    data = _read_json(path)
    return jsonify({'success': True, 'session_id': session_id, 'status': job['status'], 'data': data})


@app.route('/review/<session_id>', methods=['POST'])
def save_review(session_id):
    ctx = _require_auth()
    payload = request.get_json(silent=True) or {}
    review_data = payload.get('data') or payload.get('field_values') or {}
    import db
    job = db.get_job_by_session(session_id, ctx['tenant_id'], ctx['user_id'])
    if not job:
        return jsonify({'success': False, 'error_message': 'Session not found'}), 404
    reviewed_path = _reviewed_json_path(ctx['tenant_id'], ctx['user_id'], session_id)
    _write_json(reviewed_path, review_data)
    db.update_checklist_job(job['job_id'], ctx['tenant_id'], ctx['user_id'], status='review_ready', reviewed_data_path=str(reviewed_path), current_step='Review saved')
    return jsonify({'success': True, 'session_id': session_id, 'status': 'review_ready', 'message': 'Reviewed data saved successfully.'})


@app.route('/generate', methods=['POST'])
@app.route('/process', methods=['POST'])
def process_documents():
    ctx = _require_auth()
    tenant_id = ctx['tenant_id']; user_id = ctx['user_id']
    try:
        payload = request.get_json(silent=True) or {}
        job_id = payload.get('job_id')
        session_id = payload.get('session_id')
        selected_template_keys = payload.get('selected_templates') or []
        if not session_id or not job_id:
            return jsonify({'success': False, 'error_message': 'Missing session_id or job_id'}), 400
        if not selected_template_keys:
            return jsonify({'success': False, 'error_message': 'No templates selected'}), 400
        import db
        job = db.get_checklist_job(job_id, tenant_id, user_id)
        if not job or job.get('session_id') != session_id:
            return jsonify({'success': False, 'error_message': 'Job not found'}), 404
        if job['status'] not in ('extract_completed', 'review_ready', 'generated', 'generation_failed'):
            return jsonify({'success': False, 'error_message': 'Review is not ready yet'}), 409
        db.update_checklist_job(job_id, tenant_id, user_id, status='generating', current_step='Generating checklists')

        field_values = payload.get('field_values') or {}
        if not field_values:
            reviewed_path = Path(job.get('reviewed_data_path') or _reviewed_json_path(tenant_id, user_id, session_id))
            if reviewed_path.exists():
                field_values = _read_json(reviewed_path)
            else:
                field_values = _read_json(Path(job.get('extracted_data_path') or _extracted_json_path(tenant_id, user_id, session_id)))
        # flatten API extracted structure if needed
        flat_field_values = {}
        for k, v in (field_values or {}).items():
            if isinstance(v, dict):
                flat_field_values[k] = v.get('value', '')
            else:
                flat_field_values[k] = v
        reviewed_path = _reviewed_json_path(tenant_id, user_id, session_id)
        _write_json(reviewed_path, flat_field_values)

        output_folder = _output_folder(tenant_id, user_id, session_id)
        if output_folder.exists():
            shutil.rmtree(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)

        full_name = (flat_field_values.get('Candidate Name') or '').strip()
        if not full_name:
            full_name = 'Candidate'
        name_parts = full_name.split()
        first_name = name_parts[0] if name_parts else ''
        last_name = name_parts[-1] if len(name_parts) > 1 else ''
        role = (flat_field_values.get('Role') or '').strip()
        yes_na = 'NA' if role.upper() == 'HCA' else 'YES'
        todays = _dt.date.today().strftime('%d/%m/%Y')
        nationality = (flat_field_values.get('Nationality') or '').strip().lower()
        if 'british' in nationality or nationality == 'uk':
            for k in ['RTW Status', 'Visa Expiry Date', 'Visa Type', 'Restriction', 'Share Code']:
                flat_field_values[k] = 'UK Passport'
        replacements = {
            'Candidate Name': full_name, 'Candidate First Name': first_name, 'Candidate First name': first_name,
            'Candidate surname': last_name, 'Candidate last name': last_name,
            'Candidate Address': flat_field_values.get('Address', ''), 'Candidate address': flat_field_values.get('Address', ''),
            'Address': flat_field_values.get('Address', ''), 'Candidate Mobile Number': flat_field_values.get('Phone', ''),
            'Phone': flat_field_values.get('Phone', ''), 'DOB': flat_field_values.get('DOB', ''), 'D.O.B': flat_field_values.get('DOB', ''),
            'Date Of Birth': flat_field_values.get('DOB', ''), 'Nationality': flat_field_values.get('Nationality', ''),
            'HCA/RGN': role, ' HCA/RGN': role, 'NI Number': flat_field_values.get('NI Number', ''),
            'NMC Pin Number': flat_field_values.get('NMC PIN', ''), 'NMC Pin number': flat_field_values.get('NMC PIN', ''),
            'NMC pin': flat_field_values.get('NMC PIN', ''), 'DBS Certificate Number': flat_field_values.get('DBS Number', ''),
            'DBS certificate number': flat_field_values.get('DBS Number', ''), 'DBS Certificate issue date': flat_field_values.get('DBS Issue Date', ''),
            'DBS certificate issue date': flat_field_values.get('DBS Issue Date', ''), 'DBS Certificate last checked date': flat_field_values.get('DBS Last Checked Date', ''),
            'DBS last checked date': flat_field_values.get('DBS Last Checked Date', ''), 'DBS expiry date': flat_field_values.get('DBS Last Checked Date', ''),
            'Training completion date': flat_field_values.get('Training Date', ''), 'Training expiry date': flat_field_values.get('Training Expiry Date', ''),
            'Right to work expiry date': flat_field_values.get('Visa Expiry Date', ''), 'Right To Work Expiry Date': flat_field_values.get('Visa Expiry Date', ''),
            'Type of visa': flat_field_values.get('Visa Type', ''), 'Restriction': flat_field_values.get('Restriction', ''),
            "Today's Date": todays, "Today's date": todays, 'YES/NA': yes_na, 'Candidate share code': flat_field_values.get('Share Code', ''),
            'Form Completed By': flat_field_values.get('Form Completed By', ''), 'form completed by': flat_field_values.get('Form Completed By', ''),
            'Signature': flat_field_values.get('Signature', ''), 'signature': flat_field_values.get('Signature', ''),
            'Position': flat_field_values.get('Position', ''), 'position': flat_field_values.get('Position', ''),
        }
        for k, v in flat_field_values.items():
            replacements.setdefault(k, v or '')

        generated_files = []
        file_downloads = []
        output_rows = []
        errors = []
        for template_key in selected_template_keys:
            tmpl = get_template_by_key(template_key)
            if not tmpl:
                errors.append(f'Invalid template: {template_key}')
                continue
            template_path = TEMPLATE_FOLDER / tmpl['file_name']
            if not template_path.exists():
                errors.append(f"Template not found: {tmpl['file_name']}")
                continue
            output_filename = _build_output_filename(full_name, tmpl['template_name'], output_folder)
            output_path = output_folder / output_filename
            try:
                _, warnings = fill_docx_template(template_path, output_path, replacements)
                if warnings:
                    errors.extend([f"{tmpl['template_name']}: {w}" for w in warnings])
                generated_files.append(output_filename)
                file_downloads.append({'filename': output_filename, 'download_url': f'/download/{session_id}/{output_filename}', 'billed': False, 'template_key': tmpl['template_key'], 'template_name': tmpl['template_name']})
                output_rows.append({'output_id': f"out_{uuid.uuid4().hex}", 'template_key': tmpl['template_key'], 'template_name': tmpl['template_name'], 'output_filename': output_filename, 'output_path': str(output_path)})
            except Exception as e:
                errors.append(f"{tmpl['template_name']}: {str(e)}")
        zip_name = ''
        zip_url = ''
        if generated_files:
            zip_name = _build_zip_filename(full_name)
            zip_path = output_folder / zip_name
            with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
                for fn in generated_files:
                    zipf.write(output_folder / fn, fn)
            zip_url = f'/download-zip/{session_id}'
        if output_rows:
            db.save_output_rows(job_id=job_id, session_id=session_id, tenant_id=tenant_id, user_id=user_id, outputs=output_rows)
        db.mark_noncurrent_outputs_inactive(session_id=session_id, tenant_id=tenant_id, user_id=user_id, active_template_keys=selected_template_keys)
        db.update_checklist_job(job_id, tenant_id, user_id, status='generated' if output_rows else 'generation_failed', current_step='Generation completed' if output_rows else 'Generation failed', reviewed_data_path=str(reviewed_path), output_path=str(output_folder), generated_file_count=len(output_rows), generated_at=_dt.datetime.utcnow(), error_message='; '.join(errors[:20]) if errors else None)
        return jsonify({'success': bool(output_rows), 'job_id': job_id, 'session_id': session_id, 'status': 'generated' if output_rows else 'generation_failed', 'generated_file_count': len(output_rows), 'files': file_downloads, 'zip': {'filename': zip_name, 'download_url': zip_url} if zip_url else None, 'errors': errors})
    except Exception as e:
        app.logger.exception('generation failed')
        try:
            import db
            if job_id:
                db.update_checklist_job(job_id, tenant_id, user_id, status='generation_failed', current_step='Generation failed', error_message=str(e))
        except Exception:
            pass
        return jsonify({'success': False, 'error_message': str(e), 'status': 'generation_failed'}), 500


@app.route('/outputs/<session_id>', methods=['GET'])
def outputs_api(session_id):
    ctx = _require_auth()
    import db
    outputs = db.list_outputs(session_id, ctx['tenant_id'], ctx['user_id'])
    job = db.get_job_by_session(session_id, ctx['tenant_id'], ctx['user_id'])
    zip_url = f'/download-zip/{session_id}' if outputs else ''
    zip_name = _build_zip_filename(((_read_json(Path(job.get('reviewed_data_path') or '')) or {}).get('Candidate Name') if job and job.get('reviewed_data_path') else 'Candidate')) if outputs else ''
    files = [{'template_key': o.get('template_key'), 'template_name': o.get('template_name'), 'filename': o.get('output_filename'), 'download_url': f"/download/{session_id}/{o.get('output_filename')}", 'download_token_charged': bool(o.get('download_token_charged'))} for o in outputs]
    return jsonify({'success': True, 'session_id': session_id, 'status': job.get('status') if job else 'generated', 'files': files, 'zip': {'filename': zip_name, 'download_url': zip_url} if zip_url else None})


def _charge_for_single_output(ctx, output_row):
    import db
    if output_row.get('download_token_charged'):
        db.mark_output_downloaded(output_row['output_id'], ctx['tenant_id'], ctx['user_id'], charge_now=False)
        return True
    if db.get_tenant_tokens_remaining(ctx['tenant_id']) == 0:
        return False
    legacy_job_id = db.create_legacy_job_record(tenant_id=ctx['tenant_id'], user_id=ctx['user_id'], total_items=1, status='running')
    usage_id = db.record_usage(tenant_id=ctx['tenant_id'], user_id=ctx['user_id'], db_job_id=legacy_job_id, successful_outputs=1)
    if legacy_job_id:
        db.update_legacy_job_status(db_job_id=legacy_job_id, status='completed', successful_items=1, failed_items=0)
    db.mark_output_downloaded(output_row['output_id'], ctx['tenant_id'], ctx['user_id'], charge_usage_id=usage_id, charge_now=True)
    return True


@app.route('/download/<session_id>/<filename>')
def download(session_id, filename):
    ctx = _require_auth()
    tenant_id = ctx['tenant_id']; user_id = ctx['user_id']
    safe_name = Path(filename).name
    import db
    output_row = db.get_output_by_filename(session_id, tenant_id, user_id, safe_name)
    if not output_row:
        return jsonify({'success': False, 'error_message': 'File not found or access denied.'}), 404
    file_path = Path(output_row['output_path'])
    if not file_path.exists():
        return jsonify({'success': False, 'error_message': 'File not found or expired.'}), 404
    if not _charge_for_single_output(ctx, output_row):
        return jsonify({'success': False, 'error_message': 'No tokens remaining. Please contact NextStep to top up.'}), 402
    return send_file(file_path, as_attachment=True, download_name=safe_name)


@app.route('/download-zip/<session_id>')
def download_zip(session_id):
    ctx = _require_auth()
    import db
    outputs = db.list_outputs(session_id, ctx['tenant_id'], ctx['user_id'])
    if not outputs:
        return jsonify({'success': False, 'error_message': 'No generated files available for download.'}), 404
    job = db.get_job_by_session(session_id, ctx['tenant_id'], ctx['user_id'])
    output_folder = _output_folder(ctx['tenant_id'], ctx['user_id'], session_id)
    reviewed = _read_json(_reviewed_json_path(ctx['tenant_id'], ctx['user_id'], session_id))
    zip_name = _build_zip_filename(reviewed.get('Candidate Name') or 'Candidate')
    zip_path = output_folder / zip_name
    unbilled = [o for o in outputs if not o.get('download_token_charged')]
    need = len(unbilled)
    if need > 0:
        tokens = db.get_tenant_tokens_remaining(ctx['tenant_id'])
        if 0 <= tokens < need:
            return jsonify({'success': False, 'error_message': f'Not enough tokens for ZIP download. Need {need} token(s).'}), 402
        legacy_job_id = db.create_legacy_job_record(tenant_id=ctx['tenant_id'], user_id=ctx['user_id'], total_items=need or 1, status='running')
        usage_id = db.record_usage(tenant_id=ctx['tenant_id'], user_id=ctx['user_id'], db_job_id=legacy_job_id, successful_outputs=need) if need else None
        if legacy_job_id:
            db.update_legacy_job_status(db_job_id=legacy_job_id, status='completed', successful_items=need, failed_items=0)
        for out in outputs:
            db.mark_output_downloaded(out['output_id'], ctx['tenant_id'], ctx['user_id'], charge_usage_id=usage_id if out in unbilled else None, charge_now=out in unbilled)
    else:
        for out in outputs:
            db.mark_output_downloaded(out['output_id'], ctx['tenant_id'], ctx['user_id'], charge_now=False)
    if not zip_path.exists():
        with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
            for out in outputs:
                fp = Path(out['output_path'])
                if fp.exists():
                    zipf.write(fp, arcname=out['output_filename'])
    return send_file(zip_path, as_attachment=True, download_name=zip_name)


@app.route('/health')
def health_check():
    return jsonify({'status': 'healthy', 'gemini_configured': bool(gemini_client), 'tesseract_available': bool(pytesseract), 'pdfplumber_available': bool(pdfplumber), 'pymupdf_available': bool(fitz), 'model_fast': GEMINI_MODEL_FAST, 'model_strong': GEMINI_MODEL_STRONG, 'db': bool(os.getenv('DATABASE_URL')), 'checklist_global_concurrency': CHECKLIST_GLOBAL_CONCURRENCY})


if __name__ == '__main__':
    print('='*60)
    print('Compliance Checklist Automation - Web App')
    print('Queued extraction + isolated checklist generation')
    print('='*60)
    print(f'Gemini API configured: {bool(gemini_client)}')
    print(f'Database configured: {bool(os.getenv("DATABASE_URL"))}')
    _ensure_dispatcher_started()
    app.run(debug=True, host='0.0.0.0', port=5000)
